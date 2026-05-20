from pathlib import Path
import re
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("C:/xampp/htdocs")
excel_candidates = list(ROOT.rglob("Final*.xlsx"))
excel_path = next((p for p in excel_candidates if p.name == "Final Các hệ thống (1).xlsx"), excel_candidates[0])
project_root = excel_path.parents[1]
app_root = project_root / "Final websiteapp 1"
content_dir = excel_path.parent
out_dir = content_dir / "generated_report_assets"
out_dir.mkdir(exist_ok=True)
out_docx = content_dir / "Cafe_Connect_Final_Report_Main_Part.docx"

schema_text = (app_root / "database" / "cafe_connect_schema.sql").read_text(encoding="utf-8", errors="ignore")
tables = re.findall(r"^CREATE TABLE\s+([a-zA-Z_][a-zA-Z0-9_]*)", schema_text, flags=re.M)

route_sections = {
    "Website": ["/", "/menu", "/account", "/checkout", "/member"],
    "POS": [
        "/pos/login", "/pos/checkout", "/pos/orders", "/pos/kitchen", "/pos/customers",
        "/pos/campaigns", "/pos/inventory", "/pos/reports", "/pos/products", "/pos/staff", "/pos/cash",
    ],
    "API": [
        "member-login", "member-register", "member-lookup", "customer-create", "checkout",
        "create-order", "update-order-item", "pos-session-login", "pos-session-heartbeat",
        "pos-session-logout", "pos-session-report",
    ],
}

wb = openpyxl.load_workbook(excel_path, data_only=False)
ws = wb.active
headers = [str(c.value or "").strip() for c in next(ws.iter_rows(min_row=1, max_row=1))]
excel_rows = []
for row in ws.iter_rows(min_row=2):
    values = [c.value for c in row]
    if any(v is not None and str(v).strip() for v in values):
        excel_rows.append({headers[i] if i < len(headers) else f"col{i}": values[i] for i in range(len(values))})


FONT = "C:/Windows/Fonts/arial.ttf"
BOLD = "C:/Windows/Fonts/arialbd.ttf"
font = ImageFont.truetype(FONT, 22)
font_small = ImageFont.truetype(FONT, 18)
font_bold = ImageFont.truetype(BOLD, 24)
font_title = ImageFont.truetype(BOLD, 30)


def wrap(draw, text, font_obj, width):
    words = str(text).split()
    lines, line = [], ""
    for word in words:
        test = (line + " " + word).strip()
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= width:
            line = test
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def box(draw, xy, text, fill="#fff7ed", outline="#6f4a32", font_obj=None, radius=16):
    font_obj = font_obj or font
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    lines = wrap(draw, text, font_obj, x2 - x1 - 28)
    line_h = font_obj.size + 5
    total_h = len(lines) * line_h
    y = y1 + max(8, ((y2 - y1) - total_h) // 2)
    for line in lines:
        w = draw.textbbox((0, 0), line, font=font_obj)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, fill="#2f1b14", font=font_obj)
        y += line_h


def arrow(draw, start, end, color="#6f4a32"):
    import math
    draw.line([start, end], fill=color, width=3)
    x1, y1 = start
    x2, y2 = end
    ang = math.atan2(y2 - y1, x2 - x1)
    size = 12
    pts = [
        (x2, y2),
        (x2 - size * math.cos(ang - 0.45), y2 - size * math.sin(ang - 0.45)),
        (x2 - size * math.cos(ang + 0.45), y2 - size * math.sin(ang + 0.45)),
    ]
    draw.polygon(pts, fill=color)


def make_canvas(name, title, boxes, arrows, size=(1500, 900)):
    img = Image.new("RGB", size, "#f7f1ea")
    d = ImageDraw.Draw(img)
    d.text((40, 30), title, fill="#2f1b14", font=font_title)
    for item in boxes:
        box(d, item["xy"], item["text"], item.get("fill", "#fff7ed"), item.get("outline", "#6f4a32"), item.get("font", font))
    for item in arrows:
        arrow(d, item[0], item[1], item[2] if len(item) > 2 else "#6f4a32")
    path = out_dir / f"{name}.png"
    img.save(path)
    return path


def build_diagrams():
    org = make_canvas("org_chart", "Sơ đồ cấu trúc tổ chức Cafe Connect", [
        {"xy": (560, 100, 940, 180), "text": "Chủ chuỗi / Ban giám đốc", "fill": "#e7c7a8", "font": font_bold},
        {"xy": (170, 290, 470, 370), "text": "Quản lý vận hành", "fill": "#fff"},
        {"xy": (600, 290, 900, 370), "text": "Marketing / CRM", "fill": "#fff"},
        {"xy": (1030, 290, 1330, 370), "text": "IT / Admin hệ thống", "fill": "#fff"},
        {"xy": (120, 500, 410, 580), "text": "Quản lý cửa hàng"},
        {"xy": (450, 500, 690, 580), "text": "Thu ngân"},
        {"xy": (730, 500, 970, 580), "text": "Phục vụ"},
        {"xy": (1010, 500, 1250, 580), "text": "Barista"},
        {"xy": (520, 700, 980, 780), "text": "Khách hàng thành viên / khách lẻ", "fill": "#dcefe4"},
    ], [
        ((750, 180), (320, 290)), ((750, 180), (750, 290)), ((750, 180), (1180, 290)),
        ((320, 370), (265, 500)), ((320, 370), (570, 500)), ((750, 370), (850, 500)),
        ((1180, 370), (1130, 500)), ((570, 580), (750, 700)), ((850, 580), (750, 700)),
    ])
    current = make_canvas("current_process", "Current State Process Model", [
        {"xy": (70, 180, 300, 270), "text": "Khách mua tại quầy"},
        {"xy": (390, 180, 650, 270), "text": "Thu ngân ghi bill rời rạc"},
        {"xy": (740, 180, 1010, 270), "text": "Dữ liệu khách lưu Excel / giấy"},
        {"xy": (1100, 180, 1400, 270), "text": "Marketing gửi đại trà"},
        {"xy": (390, 430, 650, 520), "text": "Không nhận diện khách quen"},
        {"xy": (740, 430, 1010, 520), "text": "Không đo hiệu quả voucher"},
        {"xy": (1100, 430, 1400, 520), "text": "Báo cáo thủ công chậm"},
    ], [
        ((300, 225), (390, 225)), ((650, 225), (740, 225)), ((1010, 225), (1100, 225)),
        ((520, 270), (520, 430)), ((875, 270), (875, 430)), ((1250, 270), (1250, 430)),
    ])
    future = make_canvas("future_process", "Future State Process Model", [
        {"xy": (50, 170, 260, 260), "text": "Website / POS"},
        {"xy": (330, 170, 570, 260), "text": "Tra SĐT, login member"},
        {"xy": (640, 170, 900, 260), "text": "Checkout, invoice, payment"},
        {"xy": (970, 170, 1230, 260), "text": "Điểm, hạng, voucher"},
        {"xy": (1290, 170, 1480, 260), "text": "Dashboard"},
        {"xy": (330, 420, 570, 510), "text": "Customer 360"},
        {"xy": (640, 420, 900, 510), "text": "Kitchen / service order"},
        {"xy": (970, 420, 1230, 510), "text": "Campaign tracking"},
        {"xy": (1290, 420, 1480, 510), "text": "POS session report"},
    ], [
        ((260, 215), (330, 215)), ((570, 215), (640, 215)), ((900, 215), (970, 215)),
        ((1230, 215), (1290, 215)), ((450, 260), (450, 420)), ((770, 260), (770, 420)),
        ((1100, 260), (1100, 420)), ((1385, 260), (1385, 420)), ((570, 465), (640, 465)),
        ((900, 465), (970, 465)), ((1230, 465), (1290, 465)),
    ])
    usecase = make_canvas("use_case", "Use Case Diagram - Cafe Connect CRM + POS", [
        {"xy": (600, 110, 1080, 730), "text": "Cafe Connect CRM + POS", "fill": "#fffaf3", "font": font_bold},
        {"xy": (80, 130, 280, 210), "text": "Khách hàng"},
        {"xy": (80, 330, 280, 410), "text": "Thu ngân"},
        {"xy": (80, 530, 280, 610), "text": "Phục vụ / Barista"},
        {"xy": (1210, 130, 1420, 210), "text": "Marketing"},
        {"xy": (1210, 330, 1420, 410), "text": "Manager/Admin"},
        {"xy": (660, 170, 1020, 230), "text": "Đăng ký / đăng nhập member"},
        {"xy": (660, 250, 1020, 310), "text": "Tra cứu điểm, voucher, lịch sử"},
        {"xy": (660, 330, 1020, 390), "text": "Checkout & tích điểm"},
        {"xy": (660, 410, 1020, 470), "text": "Tạo service order / kitchen"},
        {"xy": (660, 490, 1020, 550), "text": "Tạo campaign / voucher"},
        {"xy": (660, 570, 1020, 630), "text": "Xem dashboard, báo cáo ca"},
    ], [
        ((280, 170), (660, 200)), ((280, 370), (660, 360)), ((280, 570), (660, 440)),
        ((1210, 170), (1020, 520)), ((1210, 370), (1020, 600)), ((280, 170), (660, 280)),
        ((1210, 370), (1020, 360)),
    ])
    erd = make_canvas("erd_overview", "Entity Relationship Overview", [
        {"xy": (60, 130, 300, 230), "text": "customers\nmembership_tiers\nsegments"},
        {"xy": (380, 130, 620, 230), "text": "vouchers\npromotions\nmarketing_emails"},
        {"xy": (700, 130, 940, 230), "text": "invoices\ninvoice_details\npayments"},
        {"xy": (1020, 130, 1260, 230), "text": "products\nproduct_images\ncategories"},
        {"xy": (60, 430, 300, 530), "text": "branches\nstaff\nstaff_shifts"},
        {"xy": (380, 430, 620, 530), "text": "pos_sessions\npos_activity_logs"},
        {"xy": (700, 430, 940, 530), "text": "service_orders\nservice_order_items\ndining_tables"},
        {"xy": (1020, 430, 1260, 530), "text": "inventory\nstock_movements\ncash_transactions"},
        {"xy": (460, 690, 860, 780), "text": "CRM insights: loyalty, campaign, performance reports", "fill": "#dcefe4", "font": font_bold},
    ], [
        ((300, 180), (380, 180)), ((620, 180), (700, 180)), ((940, 180), (1020, 180)),
        ((180, 230), (180, 430)), ((500, 230), (500, 430)), ((820, 230), (820, 430)),
        ((1140, 230), (1140, 430)), ((620, 480), (700, 480)), ((820, 530), (700, 690)),
        ((500, 530), (620, 690)), ((1140, 530), (860, 690)), ((180, 530), (460, 690)),
    ])
    ssd1 = make_canvas("ssd_member_checkout", "SSD 1 - Đăng ký thành viên và tích điểm", [
        {"xy": (60, 140, 250, 220), "text": "Customer"},
        {"xy": (330, 140, 540, 220), "text": "Website/POS UI"},
        {"xy": (620, 140, 830, 220), "text": "Auth/Customer Controller"},
        {"xy": (910, 140, 1120, 220), "text": "Invoice Model"},
        {"xy": (1200, 140, 1430, 220), "text": "MySQL DB"},
        {"xy": (330, 360, 540, 430), "text": "register/login bằng SĐT"},
        {"xy": (620, 360, 830, 430), "text": "create/find customer"},
        {"xy": (910, 360, 1120, 430), "text": "checkout + earn points"},
        {"xy": (1200, 360, 1430, 430), "text": "save customer, invoice, loyalty"},
    ], [
        ((250, 180), (330, 180)), ((540, 180), (620, 180)), ((830, 180), (910, 180)),
        ((1120, 180), (1200, 180)), ((435, 220), (435, 360)), ((725, 220), (725, 360)),
        ((1015, 220), (1015, 360)), ((1315, 220), (1315, 360)),
    ])
    ssd2 = make_canvas("ssd_voucher", "SSD 2 - Áp dụng voucher và thanh toán", [
        {"xy": (60, 140, 250, 220), "text": "Cashier / Customer"},
        {"xy": (330, 140, 540, 220), "text": "POS/Checkout UI"},
        {"xy": (620, 140, 830, 220), "text": "Voucher Model"},
        {"xy": (910, 140, 1120, 220), "text": "Invoice Model"},
        {"xy": (1200, 140, 1430, 220), "text": "MySQL DB"},
        {"xy": (330, 360, 540, 430), "text": "chọn voucher"},
        {"xy": (620, 360, 830, 430), "text": "validate active, expiry, customer"},
        {"xy": (910, 360, 1120, 430), "text": "calculate discount, paid_at"},
        {"xy": (1200, 360, 1430, 430), "text": "redeem voucher, save payment"},
    ], [
        ((250, 180), (330, 180)), ((540, 180), (620, 180)), ((830, 180), (910, 180)),
        ((1120, 180), (1200, 180)), ((435, 220), (435, 360)), ((725, 220), (725, 360)),
        ((1015, 220), (1015, 360)), ((1315, 220), (1315, 360)),
    ])
    state = make_canvas("state_machine", "State Machine - Voucher, Order, POS Session", [
        {"xy": (70, 140, 260, 210), "text": "Voucher issued"},
        {"xy": (350, 140, 540, 210), "text": "Active"},
        {"xy": (630, 100, 820, 170), "text": "Redeemed"},
        {"xy": (630, 210, 820, 280), "text": "Expired"},
        {"xy": (70, 420, 260, 490), "text": "Order waiting"},
        {"xy": (350, 420, 540, 490), "text": "Preparing"},
        {"xy": (630, 420, 820, 490), "text": "Ready"},
        {"xy": (910, 420, 1100, 490), "text": "Served/Paid"},
        {"xy": (70, 670, 260, 740), "text": "Session open"},
        {"xy": (350, 620, 540, 690), "text": "Heartbeat active"},
        {"xy": (630, 620, 820, 690), "text": "Closed manual"},
        {"xy": (630, 730, 820, 800), "text": "Closed timeout"},
    ], [
        ((260, 175), (350, 175)), ((540, 175), (630, 135)), ((540, 175), (630, 245)),
        ((260, 455), (350, 455)), ((540, 455), (630, 455)), ((820, 455), (910, 455)),
        ((260, 705), (350, 655)), ((540, 655), (630, 655)), ((540, 655), (630, 765)),
    ])
    return org, current, future, usecase, erd, ssd1, ssd2, state


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else "")
    run.font.size = Pt(9)
    run.bold = bold


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E7C7A8")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Number")


def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Title", 22, "2F1B14"),
        ("Heading 1", 16, "2F1B14"),
        ("Heading 2", 13, "6F4A32"),
        ("Heading 3", 11, "6F4A32"),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
    return doc


def build_report():
    org_img, current_img, future_img, usecase_img, erd_img, ssd1_img, ssd2_img, state_img = build_diagrams()
    doc = setup_doc()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO FINAL - CAFE CONNECT CRM + POS")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(47, 27, 20)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Xây dựng hệ thống chăm sóc khách hàng cho chuỗi 5 quán cà phê tại Hà Nội\n").bold = True
    p2.add_run("Nền tảng triển khai: PHP MVC thuần, XAMPP/Apache/MySQL, Website + POS + CRM")
    add_table(doc, ["Thông tin", "Nội dung"], [
        ["Tên hệ thống", "Cafe Connect CRM + POS"],
        ["Phạm vi", "Website khách hàng, POS vận hành, CRM/member, loyalty, campaign, inventory, report, POS session"],
        ["Cơ sở dữ liệu", "MySQL database cafe_connect_crm, schema trong database/cafe_connect_schema.sql"],
        ["Ứng dụng", "PHP MVC trong Final websiteapp 1, chạy trực tiếp bằng XAMPP"],
        ["File yêu cầu đầu vào", str(excel_path)],
        ["File xuất ra", str(out_docx)],
    ])

    doc.add_heading("Bảng đối chiếu task từ Excel", level=1)
    rows = []
    for item in excel_rows:
        rows.append([
            item.get("Main part", "") or "",
            item.get("Task", "") or "",
            item.get("PiC", "") or "",
            item.get("Status", "") or "",
            "Đã hoàn thiện trong báo cáo này",
        ])
    add_table(doc, ["Main part", "Task", "PiC", "Status Excel", "Outcome cập nhật"], rows)

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Cafe Connect CRM + POS là hệ thống thông tin được thiết kế cho chuỗi 5 quán cà phê tại Hà Nội. "
        "Dự án giải quyết bài toán dữ liệu khách hàng phân tán giữa quầy bán hàng, website, marketing và báo cáo quản lý. "
        "Phiên bản hiện tại đã được xây dựng bằng PHP MVC thuần trên XAMPP, sử dụng MySQL để lưu dữ liệu thật và có các route riêng "
        "cho website khách hàng, POS thu ngân/phục vụ/pha chế, CRM, campaign, inventory, report và phiên làm việc POS."
    )
    doc.add_paragraph(
        "Mục tiêu của tài liệu là hoàn thiện toàn bộ phần Main part trong file kế hoạch: thu thập yêu cầu, phân tích quy trình, "
        "mô hình hóa hệ thống, thiết kế nâng cao, database/query, trigger, kiểm thử và thảo luận về AI trong kinh doanh."
    )

    doc.add_heading("Part A: Requirement Gathering and Analysis", level=1)
    doc.add_heading("I. Initial Project Analysis", level=2)
    doc.add_heading("1.1. Bối cảnh kinh doanh", level=3)
    doc.add_paragraph(
        "Cafe Connect là chuỗi 5 cửa hàng cà phê tại Hà Nội, phục vụ khách mua tại quầy, khách ngồi tại bàn và khách đặt hàng qua website. "
        "Mô hình kinh doanh kết hợp bán đồ uống, đồ ăn nhẹ, membership, voucher/campaign và chăm sóc khách hàng sau bán. "
        "Trước khi số hóa, dữ liệu khách hàng, hóa đơn, điểm tích lũy và hiệu quả marketing dễ bị phân tán, khiến việc cá nhân hóa ưu đãi "
        "và đánh giá hiệu quả vận hành chậm."
    )
    doc.add_heading("1.2. Các bên liên quan chính", level=3)
    add_table(doc, ["Stakeholder", "Mối quan tâm", "Tương tác hệ thống"], [
        ["Chủ chuỗi / Ban giám đốc", "Doanh thu, tăng trưởng khách quay lại, hiệu quả chi nhánh", "Dashboard, báo cáo ca, doanh thu, chiến dịch"],
        ["Quản lý cửa hàng", "Điều phối nhân sự, tồn kho, hiệu suất ca", "POS reports, inventory, staff, cash"],
        ["Marketing / CRM", "Phân nhóm khách, phát hành voucher, đo campaign", "Campaigns, vouchers, newsletter, member data"],
        ["Thu ngân", "Checkout nhanh, tra khách, áp voucher, tạo khách mới", "POS checkout, customer-create, payment"],
        ["Phục vụ", "Tạo order bàn và gửi xuống bếp", "POS orders, dining tables"],
        ["Barista", "Theo dõi queue và cập nhật món", "POS kitchen, service_order_items"],
        ["Khách hàng", "Đăng ký thành viên, xem điểm/voucher, đặt hàng", "Website account, member, menu, checkout"],
        ["IT/Admin", "Cấu hình hệ thống, phân quyền, database", "install.php, staff, schema, API"],
    ])
    doc.add_heading("1.3. Mục tiêu hệ thống", level=3)
    add_bullets(doc, [
        "Số hóa hồ sơ khách hàng tập trung theo số điện thoại/email.",
        "Kết nối dữ liệu POS và website vào cùng một customer profile.",
        "Tự động cộng điểm, nâng hạng thành viên và quản lý voucher.",
        "Tối ưu vận hành tại quầy: order bàn, kitchen queue, checkout, thu chi, tồn kho.",
        "Cung cấp dashboard doanh thu, top sản phẩm, hiệu quả campaign và hiệu suất nhân viên theo ca.",
        "Giảm thao tác thủ công và tạo nền tảng cho phân tích dữ liệu/AI trong tương lai.",
    ])
    doc.add_heading("1.4. Sơ đồ cấu trúc tổ chức", level=3)
    doc.add_picture(str(org_img), width=Inches(6.5))

    doc.add_heading("II. Functional Requirements", level=2)
    add_table(doc, ["Module", "Yêu cầu chức năng", "Tiêu chí nghiệm thu"], [
        ["Member/Auth website", "Đăng ký/đăng nhập bằng số điện thoại, tên, email; lưu session member; đăng xuất.", "Khách đăng ký xong xem được điểm/voucher/lịch sử và checkout bằng đúng customer_id."],
        ["Menu website", "Hiển thị sản phẩm từ bảng products/product_images/categories; thêm vào giỏ localStorage.", "/menu load dữ liệu MySQL, chuyển /checkout không mất giỏ."],
        ["Website checkout", "Tính subtotal, giảm hạng, voucher, payment method; tạo invoice website.", "Ghi invoices, invoice_details, payments, loyalty_point_transactions; cập nhật customers."],
        ["POS login/session", "Chọn nhân viên/role từ DB; tạo pos_sessions; heartbeat; logout đóng ca.", "API ghi POS thiếu token/session phải trả lỗi."],
        ["POS cashier", "Tra khách, tạo khách, chọn món, áp voucher, checkout direct/order.", "Invoice POS có bill_started_at, paid_at, pos_session_id."],
        ["Service order", "Phục vụ chọn bàn, tạo order, gửi xuống bếp.", "service_orders và service_order_items được tạo, bàn occupied."],
        ["Kitchen queue", "Barista chuyển waiting/preparing/ready/served.", "Lưu preparing_started_at, ready_at, served_at, prepared_by_session_id."],
        ["Campaign/voucher", "Marketing tạo promotion, phát hành voucher theo nhóm khách.", "Đo issued, redeemed, attributed revenue."],
        ["Inventory/cash", "Quản lý nhập/xuất kho và thu/chi quầy.", "stock_movements/cash_transactions gắn session."],
        ["Dashboard/report", "Báo cáo doanh thu, top sản phẩm, tồn thấp, campaign, session performance.", "/pos/reports hiển thị phiên, doanh thu, order, món pha, log."],
    ])

    doc.add_heading("III. User Stories", level=2)
    story_rows = []
    stakeholders = {
        "Customer": [("đăng ký bằng số điện thoại", "tích điểm và nhận ưu đãi cá nhân"), ("xem điểm, hạng và voucher", "biết quyền lợi trước khi mua"), ("đặt hàng trên website", "mua nhanh mà vẫn được cộng điểm")],
        "Cashier": [("tra cứu khách bằng SĐT", "áp đúng hạng/voucher"), ("checkout nhanh tại quầy", "giảm thời gian chờ"), ("tạo khách mới trên POS", "đồng bộ ngay với website member")],
        "Waiter": [("chọn bàn và tạo order", "gửi yêu cầu xuống bếp chính xác"), ("xem trạng thái bàn", "phục vụ đúng thứ tự"), ("ghi chú order", "đáp ứng yêu cầu ít đá/ít đường")],
        "Barista": [("xem kitchen queue", "biết món nào cần pha trước"), ("cập nhật trạng thái món", "phục vụ và thu ngân theo dõi được"), ("xem số món đã pha theo ca", "đánh giá hiệu suất làm việc")],
        "Marketing": [("tạo campaign/voucher", "chăm sóc đúng nhóm khách"), ("xem tỷ lệ voucher redeemed", "đánh giá hiệu quả campaign"), ("quản lý newsletter", "tăng tương tác khách hàng")],
        "Manager/Admin": [("xem dashboard doanh thu", "ra quyết định vận hành"), ("quản lý sản phẩm/nhân viên", "duy trì dữ liệu chuẩn"), ("xem báo cáo phiên POS", "kiểm soát hiệu suất nhân viên theo ca")],
    }
    criteria = "AC1 giao diện hiển thị đúng; AC2 backend validate quyền/dữ liệu; AC3 dữ liệu lưu vào MySQL; AC4 sau refresh vẫn truy xuất được qua báo cáo/API."
    for actor, items in stakeholders.items():
        for want, value in items:
            story_rows.append([actor, f"As a {actor}, I want {want}, so that {value}.", criteria])
    add_table(doc, ["User", "User story", "4 tiêu chí chấp nhận"], story_rows)

    doc.add_heading("IV. Epics", level=2)
    add_table(doc, ["Epic", "Mục tiêu", "Module hiện có"], [
        ["Customer Profile 360", "Quản lý hồ sơ, điểm, lịch sử, favorite, voucher", "Customer, Website account/member, POS customers"],
        ["Loyalty & Membership", "Tích điểm, nâng hạng, ưu đãi theo tier", "Invoice, Voucher, loyalty_point_transactions"],
        ["Marketing Automation", "Campaign, newsletter, voucher theo segment", "Campaign, Promotion, Vouchers, campaign_recipients"],
        ["Omnichannel POS & Website", "Dùng chung sản phẩm, checkout, customer profile", "Website routes, POS routes, api.php"],
        ["Operations & Session Performance", "Order bàn, kitchen, cash, kho, ca làm", "Order, Inventory, PosSession, Report"],
        ["Analytics & Reporting", "Dashboard cho doanh thu, campaign, nhân viên", "Dashboard, Report, session_reports"],
    ])

    doc.add_heading("V. Non-Functional Requirements", level=2)
    add_table(doc, ["Nhóm NFR", "Yêu cầu", "Thiết kế đáp ứng"], [
        ["Hiệu năng", "Route POS/website phản hồi nhanh cho demo XAMPP; query giới hạn dữ liệu dashboard.", "PDO, index SQL, bootstrap theo module, JS page-scoped."],
        ["Độ sẵn sàng", "Có install/reset DB để khôi phục nhanh môi trường nghiệm thu.", "install.php import cafe_connect_schema.sql."],
        ["Khả năng mở rộng", "Tách MVC, Model riêng cho Customer/Product/Invoice/Order/PosSession.", "Có thể thêm controller/model mới không phá layout."],
        ["Bảo mật demo", "Validate role/session/token cho API POS ghi dữ liệu.", "AuthController + PosSession::requireOpen."],
        ["Tương thích", "Chạy trên Apache/XAMPP; responsive desktop/mobile.", "PHP thuần, CSS responsive, fallback index.php?route=."],
        ["Dữ liệu", "Ràng buộc khóa chính/khóa ngoại, unique phone/email, status enum.", "31 bảng MySQL với FK/index."],
        ["Khả dụng vận hành", "POS topbar, role nav, report theo ca rõ ràng.", "/pos/login và /pos/reports."],
    ])

    doc.add_heading("Part B: System Analysis Modelling", level=1)
    doc.add_heading("I. Current State Process Model", level=2)
    doc.add_picture(str(current_img), width=Inches(6.5))
    doc.add_paragraph("Quy trình hiện tại giả định dữ liệu bán hàng, thông tin khách, voucher và báo cáo còn tách rời. Nhân viên khó nhận diện khách quen tại quầy, marketing không đo được voucher theo hóa đơn, quản lý phải tổng hợp thủ công từ nhiều nguồn.")
    doc.add_heading("II. Opportunities for Improvement", level=2)
    add_bullets(doc, [
        "Chuẩn hóa customer_id theo số điện thoại để tránh trùng khách.",
        "Kết nối checkout với loyalty để cộng điểm/nâng hạng tự động.",
        "Tạo voucher cá nhân theo segment thay vì khuyến mãi đại trà.",
        "Gắn hóa đơn, kitchen item, cash và stock movement với pos_session để đo hiệu suất ca.",
        "Tách role POS giúp giao diện gọn và backend kiểm soát quyền.",
        "Dashboard hóa dữ liệu để quản lý nhìn doanh thu/campaign/tồn kho gần thời gian thực.",
    ])
    doc.add_heading("III. Future State Process Model", level=2)
    doc.add_picture(str(future_img), width=Inches(6.5))
    doc.add_paragraph("Quy trình tương lai lấy POS và website làm nguồn dữ liệu nghiệp vụ. Mỗi lần khách mua hàng, hệ thống đồng thời ghi invoice, payment, loyalty transaction, voucher redemption và session performance. Dữ liệu này được tái sử dụng cho member portal, campaign và report.")
    doc.add_heading("IV. Use Case Diagram", level=2)
    doc.add_picture(str(usecase_img), width=Inches(6.5))
    add_table(doc, ["Use case", "Actor chính", "Kết quả"], [
        ["Đăng ký/đăng nhập member", "Khách hàng", "Tạo hoặc mở customer profile."],
        ["Checkout và tích điểm", "Thu ngân/Khách hàng", "Tạo invoice, payment, cộng điểm, cập nhật hạng."],
        ["Tạo service order", "Phục vụ", "Order bàn được gửi xuống bếp."],
        ["Cập nhật kitchen queue", "Barista", "Món chuyển trạng thái và log hiệu suất pha chế."],
        ["Tạo campaign/voucher", "Marketing", "Voucher được phát hành cho nhóm khách."],
        ["Xem báo cáo phiên", "Manager/Admin", "Theo dõi doanh thu, thời lượng, việc đã làm theo ca."],
    ])
    doc.add_heading("V. Entity Relationship", level=2)
    doc.add_picture(str(erd_img), width=Inches(6.5))
    doc.add_paragraph("Schema hiện tại gồm 31 bảng. Các nhóm quan trọng: " + ", ".join(tables[:31]) + ".")

    doc.add_heading("Part C: System Design and Advanced Modelling", level=1)
    doc.add_heading("I. Use Case Narratives", level=2)
    add_table(doc, ["Mục", "Use case 1: Đăng ký thành viên và tích điểm"], [
        ["User story liên quan", "As a Customer, I want đăng ký bằng SĐT, so that tôi được tích điểm và nhận ưu đãi cá nhân."],
        ["Stakeholders", "Khách hàng, thu ngân, marketing, quản lý."],
        ["Điều kiện tiên quyết", "Database đã cài; sản phẩm/tier tồn tại; khách cung cấp SĐT hợp lệ."],
        ["Tác nhân kích hoạt", "Khách đăng ký ở /account hoặc thu ngân tạo khách tại POS."],
        ["Cam kết tối thiểu", "Nếu lỗi, không tạo dữ liệu trùng SĐT; thông báo rõ lỗi."],
        ["Cam kết thành công", "Customer profile được tạo/mở, checkout tạo invoice, điểm và tổng chi tiêu cập nhật."],
        ["Luồng chính", "1. Nhập SĐT/tên/email. 2. Hệ thống kiểm tra khách cũ. 3. Tạo customer nếu chưa có. 4. Khách/thu ngân checkout. 5. Hệ thống ghi invoice/payment/detail. 6. Cộng điểm, nâng hạng nếu đủ điều kiện. 7. Member portal hiển thị lịch sử mới."],
        ["Ngoại lệ", "SĐT trùng: mở hồ sơ cũ; giỏ rỗng: không checkout; DB chưa install: yêu cầu chạy install.php."],
    ])
    add_table(doc, ["Mục", "Use case 2: Áp dụng voucher / reward redemption"], [
        ["User story liên quan", "As a Customer, I want dùng voucher hợp lệ, so that tôi được giảm giá khi thanh toán."],
        ["Stakeholders", "Khách hàng, thu ngân, marketing, quản lý."],
        ["Điều kiện tiên quyết", "Customer đã có voucher active/issued, voucher chưa hết hạn/chưa redeemed."],
        ["Tác nhân kích hoạt", "Khách chọn voucher ở website checkout hoặc thu ngân áp voucher tại POS."],
        ["Cam kết tối thiểu", "Voucher không hợp lệ không làm thay đổi invoice hoặc trạng thái voucher."],
        ["Cam kết thành công", "Invoice lưu voucher_id, discount, payment; voucher chuyển redeemed và campaign đo được attributed revenue."],
        ["Luồng chính", "1. Lookup customer. 2. Load voucher usable. 3. Chọn voucher. 4. Voucher model validate customer/expiry/status. 5. Invoice model tính giảm giá. 6. Ghi hóa đơn và payment. 7. Redeem voucher. 8. Dashboard campaign cập nhật."],
        ["Ngoại lệ", "Voucher hết hạn/đã dùng: trả lỗi; voucher không thuộc customer: từ chối; tổng tiền sau giảm không âm."],
    ])
    doc.add_heading("II. System Sequence Diagram (SSD)", level=2)
    doc.add_picture(str(ssd1_img), width=Inches(6.5))
    doc.add_picture(str(ssd2_img), width=Inches(6.5))
    doc.add_heading("III. State Machine Diagram", level=2)
    doc.add_picture(str(state_img), width=Inches(6.5))
    doc.add_paragraph("Voucher chuyển từ issued sang active khi còn hiệu lực, sau đó redeemed khi được dùng hoặc expired khi quá hạn. Service order item chuyển waiting -> preparing -> ready -> served. POS session chuyển open -> heartbeat active -> closed manual hoặc closed timeout.")

    doc.add_heading("IV. User Interface Design", level=2)
    add_table(doc, ["Thực thể/UI", "Route", "Thiết kế/Chức năng"], [
        ["Khách hàng / Member profile", "/account, /member, /pos/customers", "Form login/register bằng SĐT, profile hạng/điểm/voucher/favorite/lịch sử, POS tạo khách mới."],
        ["Sản phẩm / Menu", "/menu, /pos/checkout", "Product card từ MySQL, category, ảnh, giá, giỏ hàng localStorage, POS product picker."],
        ["Order / Checkout", "/checkout, /pos/checkout, /pos/orders", "Cart, voucher, payment method, bill_started_at, checkout direct/service order."],
        ["Kitchen item", "/pos/kitchen", "Board món chờ, nút preparing/ready/served, ghi timestamp và staff/session pha chế."],
        ["Marketing campaign", "/pos/campaigns", "Form tạo promotion, số voucher phát hành, redeemed và doanh thu liên quan."],
        ["Session report", "/pos/reports", "Bảng phiên làm việc: thời lượng, doanh thu, order, món pha, thu/chi, log chính theo ca."],
    ])

    doc.add_heading("V. Database Implementation & Queries", level=2)
    doc.add_paragraph(f"Database cafe_connect_crm hiện có {len(tables)} bảng, dùng InnoDB, utf8mb4, khóa ngoại và index cho các quan hệ chính. Các bảng lõi gồm customers, membership_tiers, branches, staff, pos_sessions, products, promotions, vouchers, service_orders, invoices, payments, invoice_details, loyalty_point_transactions, pos_activity_logs, stock_movements và cash_transactions.")
    add_table(doc, ["Nhóm bảng", "Bảng tiêu biểu", "Vai trò"], [
        ["CRM/member", "customers, membership_tiers, customer_segments, customer_favorites", "Hồ sơ khách, hạng, segment, yêu thích."],
        ["POS/order", "dining_tables, service_orders, service_order_items, invoices, invoice_details, payments", "Bán hàng, order bàn, payment."],
        ["Loyalty/campaign", "promotions, vouchers, marketing_emails, campaign_recipients, loyalty_point_transactions", "Ưu đãi, voucher, email và điểm."],
        ["Operation/session", "staff, staff_shifts, pos_sessions, pos_activity_logs", "Nhân viên, role, ca làm, log hiệu suất."],
        ["Inventory/cash", "branch_inventory, inventory_materials, stock_movements, cash_transactions", "Tồn kho, nhập/xuất, thu chi."],
    ])
    doc.add_paragraph("Các query nghiệp vụ mẫu:")
    add_code(doc, """
-- Top 10 khách hàng chi tiêu cao nhất trong tháng
SELECT c.customer_name, c.phone_number, mt.tier_name, SUM(i.total_amount) AS monthly_spending
FROM customers c
JOIN membership_tiers mt ON mt.id = c.membership_tier_id
JOIN invoices i ON i.customer_id = c.id
WHERE i.status = 'paid' AND i.invoice_date BETWEEN '2026-05-01' AND '2026-05-31'
GROUP BY c.id, c.customer_name, c.phone_number, mt.tier_name
ORDER BY monthly_spending DESC
LIMIT 10;

-- Hiệu quả chiến dịch marketing
SELECT p.promotion_name, COUNT(v.id) AS issued_vouchers,
       SUM(v.status = 'redeemed') AS redeemed_vouchers,
       ROUND(SUM(v.status = 'redeemed') / NULLIF(COUNT(v.id), 0) * 100, 2) AS redemption_rate,
       COALESCE(SUM(i.total_amount), 0) AS attributed_revenue
FROM promotions p
LEFT JOIN vouchers v ON v.promotion_id = p.id
LEFT JOIN invoices i ON i.voucher_id = v.id AND i.status = 'paid'
GROUP BY p.id, p.promotion_name;

-- Khung giờ khách thành viên thường đến
SELECT HOUR(i.paid_at) AS visit_hour, COUNT(*) AS member_invoice_count, SUM(i.total_amount) AS revenue
FROM invoices i
WHERE i.customer_id IS NOT NULL AND i.status = 'paid'
GROUP BY HOUR(i.paid_at)
ORDER BY member_invoice_count DESC;

-- Hiệu suất theo phiên POS
SELECT ps.id, s.staff_name, ps.staff_role,
       TIMESTAMPDIFF(MINUTE, ps.opened_at, COALESCE(ps.closed_at, NOW())) AS duration_minutes,
       COUNT(i.id) AS invoice_count, COALESCE(SUM(i.total_amount),0) AS revenue
FROM pos_sessions ps
JOIN staff s ON s.id = ps.staff_id
LEFT JOIN invoices i ON i.pos_session_id = ps.id AND i.status = 'paid'
GROUP BY ps.id, s.staff_name, ps.staff_role, ps.opened_at, ps.closed_at;
""")
    doc.add_heading("VI. Triggers", level=2)
    doc.add_paragraph("Trong phiên bản PHP hiện tại, logic trigger nghiệp vụ được triển khai ở tầng ứng dụng: Invoice::checkout cộng điểm, cập nhật tổng chi tiêu, gọi upgradeTier; Voucher::redeem chuyển trạng thái voucher; PosSession log ghi hoạt động. Nếu yêu cầu triển khai trigger ở MySQL, có thể dùng các trigger sau như phương án tương đương:")
    add_code(doc, """
DELIMITER $$
CREATE TRIGGER trg_after_invoice_paid
AFTER INSERT ON invoices
FOR EACH ROW
BEGIN
  IF NEW.status = 'paid' AND NEW.customer_id IS NOT NULL THEN
    UPDATE customers
    SET current_points = current_points + NEW.points_earned,
        total_spending = total_spending + NEW.total_amount,
        last_visit_date = NEW.invoice_date
    WHERE id = NEW.customer_id;

    UPDATE customers c
    JOIN membership_tiers mt ON mt.min_total_spending = (
      SELECT MAX(mt2.min_total_spending)
      FROM membership_tiers mt2
      WHERE mt2.min_total_spending <= c.total_spending
    )
    SET c.membership_tier_id = mt.id
    WHERE c.id = NEW.customer_id;
  END IF;
END$$

CREATE TRIGGER trg_after_invoice_voucher
AFTER INSERT ON invoices
FOR EACH ROW
BEGIN
  IF NEW.status = 'paid' AND NEW.voucher_id IS NOT NULL THEN
    UPDATE vouchers SET status = 'redeemed', used_at = NEW.paid_at WHERE id = NEW.voucher_id;
  END IF;
END$$
DELIMITER ;
""")

    doc.add_heading("Part D: Testing", level=1)
    add_table(doc, ["Kịch bản", "Bước kiểm thử", "Kết quả mong đợi"], [
        ["Happy path: member + checkout + voucher", "1. Chạy install.php. 2. Website đăng ký member mới. 3. Thêm sản phẩm ở /menu. 4. Checkout. 5. Mở /member.", "Member tạo thành công; invoice website có customer_id; điểm tăng; lịch sử hóa đơn hiển thị."],
        ["Happy path: POS service order", "1. Login waiter. 2. Tạo order bàn. 3. Login barista update ready. 4. Login cashier checkout order.", "Order có item; kitchen timestamp lưu; invoice POS có pos_session_id, bill_started_at từ order.created_at, paid_at."],
        ["Alternative: voucher không hợp lệ", "1. Dùng voucher đã redeemed/hết hạn. 2. Checkout.", "API trả ok=false; invoice không tạo hoặc không áp discount; voucher không đổi trạng thái."],
        ["Alternative: thiếu POS session", "Gọi create-order/update-order-item/checkout POS không gửi session_token.", "API trả lỗi quyền/session, không ghi dữ liệu."],
        ["Regression smoke API", "Chạy tests/smoke_api.ps1.", "member-register, checkout, create-order, kitchen update, dashboard, session-report, logout đều ok=true."],
    ])
    doc.add_paragraph("Kiểm thử kỹ thuật đã dùng trong dự án: php -l toàn bộ file PHP, node --check assets/js/app.js, route check HTTP 200, smoke_api.ps1, và reset database sau kiểm thử để dữ liệu mẫu sạch.")

    doc.add_heading("Part E: Discussion - Ứng dụng AI trong kinh doanh và hệ thống CRM/POS", level=1)
    doc.add_paragraph(
        "AI có giá trị thực tế lớn đối với hệ thống CRM/POS như Cafe Connect vì nó biến dữ liệu giao dịch hằng ngày thành quyết định kinh doanh nhanh hơn. "
        "Trong hệ thống hiện tại, mỗi hóa đơn, voucher, lượt truy cập website, order bàn, trạng thái pha chế và phiên làm việc POS đều được lưu có cấu trúc. "
        "Đây là nền dữ liệu cần thiết để AI dự báo khách có khả năng quay lại, gợi ý voucher phù hợp, dự đoán nhu cầu nguyên vật liệu theo khung giờ và phát hiện bất thường trong thu chi hoặc hiệu suất ca.\n\n"
        "Các nghiên cứu về AI trong dịch vụ cho thấy AI phù hợp nhất khi bắt đầu từ các tác vụ có tính cơ học và phân tích, ví dụ phân loại khách hàng, dự báo doanh thu, đề xuất sản phẩm hoặc tự động hóa báo cáo (Huang & Rust, 2018). "
        "Với Cafe Connect, điều này có nghĩa là AI không thay thế hoàn toàn nhân viên quầy, mà hỗ trợ họ ra quyết định: thu ngân biết khách nên dùng voucher nào, marketing biết nhóm khách nào cần chiến dịch tái kích hoạt, quản lý biết ca nào đang thiếu nhân sự. "
        "Quan điểm này phù hợp với nghịch lý automation-augmentation: doanh nghiệp nên kết hợp tự động hóa với tăng cường năng lực con người thay vì chỉ cắt giảm lao động (Raisch & Krakowski, 2021).\n\n"
        "Tuy nhiên, AI chỉ tạo giá trị khi dữ liệu đủ tin cậy và quy trình được quản trị rõ. Shrestha, Ben-Menahem và von Krogh (2019) nhấn mạnh quyết định tổ chức cần cấu trúc phù hợp giữa con người và thuật toán. "
        "Vì vậy, hệ thống phải lưu nguồn dữ liệu, thời điểm, người thao tác và phiên làm việc. POS session và activity logs trong dự án là nền tảng tốt cho yêu cầu này. "
        "Các tổng quan nghiên cứu cũng chỉ ra rằng AI tạo giá trị qua hỗ trợ quyết định, tự động hóa, tăng tương tác khách hàng và tạo dịch vụ mới (Borges et al., 2021; Dwivedi et al., 2021). "
        "Trong tương lai, Cafe Connect có thể bổ sung mô hình dự báo churn, phân khúc RFM tự động, gợi ý combo theo lịch sử mua và cảnh báo tồn kho. "
        "Điều kiện triển khai là bảo vệ dữ liệu cá nhân, minh bạch tiêu chí gợi ý, tránh thiên lệch với nhóm khách ít dữ liệu và cho phép quản lý kiểm tra lại quyết định AI. "
        "Như vậy, AI nên được xem là lớp phân tích nâng cao trên CRM/POS, không phải thay thế quy trình nghiệp vụ nền tảng."
    )

    doc.add_heading("Slide Outline", level=1)
    add_numbered(doc, [
        "Bối cảnh Cafe Connect và vấn đề dữ liệu phân tán",
        "Mục tiêu hệ thống CRM + POS",
        "Stakeholders và yêu cầu chính",
        "Kiến trúc PHP MVC/XAMPP/MySQL",
        "Website member/menu/checkout",
        "POS role và session theo ca",
        "ERD và các bảng lõi",
        "Use case member + voucher",
        "Dashboard/report và campaign performance",
        "Testing và smoke API",
        "AI trong CRM/POS",
        "Kết luận và hướng phát triển",
    ])
    doc.add_heading("Format Report Checklist", level=1)
    add_bullets(doc, [
        "Có trang bìa và thông tin dự án.",
        "Tách rõ Introduction, Part A, Part B, Part C, Part D, Part E.",
        "Có bảng đối chiếu toàn bộ task từ file Excel.",
        "Có sơ đồ tổ chức, process model, use case, ERD, SSD và state machine.",
        "Có database/query/trigger/testing.",
        "Có thảo luận AI khoảng 400 từ và danh mục nguồn học thuật.",
        "Dùng định dạng heading/table thống nhất để dễ xuất nộp.",
    ])
    doc.add_heading("Tài liệu tham khảo", level=1)
    refs = [
        "Borges, A. F. S., Laurindo, F. J. B., Spínola, M. M., Gonçalves, R. F., & Mattos, C. A. (2021). The strategic use of artificial intelligence in the digital era: Systematic literature review and future research directions. International Journal of Information Management, 57, 102225. https://doi.org/10.1016/j.ijinfomgt.2020.102225",
        "Dwivedi, Y. K., et al. (2021). Artificial Intelligence (AI): Multidisciplinary perspectives on emerging challenges, opportunities, and agenda for research, practice and policy. International Journal of Information Management, 57, 101994. https://doi.org/10.1016/j.ijinfomgt.2019.08.002",
        "Huang, M.-H., & Rust, R. T. (2018). Artificial Intelligence in Service. Journal of Service Research, 21(2), 155-172. https://doi.org/10.1177/1094670517752459",
        "Raisch, S., & Krakowski, S. (2021). Artificial Intelligence and Management: The Automation-Augmentation Paradox. Academy of Management Review, 46(1), 192-210. https://doi.org/10.5465/amr.2018.0072",
        "Shrestha, Y. R., Ben-Menahem, S. M., & von Krogh, G. (2019). Organizational Decision-Making Structures in the Age of Artificial Intelligence. California Management Review, 61(4), 66-83. https://doi.org/10.1177/0008125619862257",
        "Wamba-Taguimdje, S. L., Wamba, S. F., Kamdjoug, J. R. K., & Wanko, C. E. T. (2020). Influence of Artificial Intelligence (AI) on Firm Performance: The Business Value of AI-Based Transformation Projects. Business Process Management Journal, 26(7), 1893-1924. https://doi.org/10.1108/BPMJ-10-2019-0411",
        "Cafe Connect project source: README.md, cafe_connect_schema.sql, MVC controllers/models and smoke_api.ps1 in Final websiteapp 1.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")
    doc.add_heading("Phụ lục: Route triển khai hiện tại", level=1)
    for group, items in route_sections.items():
        doc.add_heading(group, level=2)
        add_bullets(doc, items)

    doc.save(out_docx)
    return out_docx, len(doc.paragraphs), len(doc.tables), len(list(out_dir.glob("*.png")))


if __name__ == "__main__":
    path, paragraphs, table_count, diagram_count = build_report()
    print(path)
    print(f"paragraphs={paragraphs} tables={table_count} diagrams={diagram_count}")
